#
#  pip3 install docx
#  pip3 install pyttsx3
#
from docx import Document
from docx.shared import Inches
from docx.shared import Cm
import pyttsx3

pyttsx3.speak('Lets start,')
pyttsx3.speak('Please fill you CV')

def speak(text):
    pyttsx3.speak(text)

name = 'John'
phone_number = '32(0).123.45.78'
email = "jj@me.com"

speak('Hello ' + name)

document =  Document()

p = document.add_paragraph()

document.add_paragraph('Name: ' + name)
document.add_paragraph('Phone number: ' + phone_number)
document.add_paragraph('email: ' + email)
document.add_paragraph('')


document.add_paragraph('')

p.add_run('Company Blablow').bold = True
p.add_run(' more text....' + '\n\n').italic = True

# more experiences
while True:
    has_more_experiences = input('Do you have more experiences (Y or N)? ')
    if has_more_experiences.lower() == 'y':

        company = input('Enter company ')
        from_date = input('From date ')
        to_date = input('To date ')
        experience = input('Describe your experience at ' + company.capitalize() + ' ? ')
        print('')

        p.add_run(company + '\n').bold = True
        p.add_run('From ' + from_date + ' to ' + to_date + '\n')
        p.add_run('Experience: ' + experience + '\n')
        p.add_run('\n')

    else:
        break

# Skills
document.add_heading('Skills:')
skill = input('Enter skill ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills (Y or N)? ')
    if has_more_skills.lower() == 'y':

        skill = input('Enter skill ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
        print('')

    else:
        break
p.add_run('\n')


p.add_run(' ____________')
p.add_run('\n')

# Pictures
document.add_heading('Pictures:')
document.add_picture('test.jpg', width=Inches(2.0))
document.add_paragraph('')
document.add_picture('test.jpg', width=Cm(4))
document.add_paragraph('')
document.add_paragraph(' ..............')
document.add_paragraph('')

# footer
section = document.sections[0]
footer = section.footer
pp = footer.paragraphs[0]
pp.text = "---  CV generated using Python  ---"

document.save("cv.docx")

print('   END.')
